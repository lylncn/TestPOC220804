import pyqtgraph as pg
from PyQt5.QtCore import QTimer, Qt, QDir
from PyQt5.QtWidgets import QFileDialog, QSizePolicy, QSpacerItem
from pyqtgraph.Qt import QtGui, QtCore, QtWidgets
import pyqtgraph.console
import numpy as np
import array
import serial
import serial.tools.list_ports
import sys
from functools import partial

from PyQt5.QtGui import QPalette, QPixmap
from MySerialClass import MySerialClass

from pyqtgraph.dockarea import *
import MyFlukeWidget
import MyTestItemTree
import docx

idx = 0


def frameHandle():
    global ser1
    global idx
    if ser1 is not None:
        if ser1.frameDataUpdate == 1:
            ser1.frameDataUpdate = 0
            tmp = np.frombuffer(ser1.frameData[2:18], dtype=np.int16)
            data[idx] = tmp
            idx += 1
            dataDisplay()


def dataDisplay():
    global idx
    global w3
    global dispState
    global isDisplay
    lmax = np.zeros(8)
    lmin = np.zeros(8)
    lmin.fill(65535.0)

    if dispState == 0:
        if idx > 0:
            if idx < historyLength:
                for n in range(0, 8):
                    #labelCurve[n].setText(str(data[idx -1, n]))
                    listTableItem[n][0].setText(str(data[idx -1, n]))
                    if isDisplay[n] is True:
                        curve[n].setData(np.arange(0, idx), data[0:idx, n])
                        lmax[n] = max(data[0:idx, n])
                        lmin[n] = min(data[0:idx, n])
                    else:
                        curve[n].setData([0])
                nmax = max(lmax)
                nmin = min(lmin)
                w3.setRange(xRange=[0, historyLength], yRange=[nmin - 10, nmax + 10], padding=0)
            else:
                for n in range(0, 8):
                    #labelCurve[n] = data[idx, n]
                    if isDisplay[n] is True:
                        curve[n].setData(np.arange(idx - historyLength, idx), data[idx - historyLength:idx, n])
                        lmax[n] = max(data[idx - historyLength:idx, n])
                        lmin[n] = min(data[idx - historyLength:idx, n])
                    else:
                        curve[n].setData([0])
                nmax = max(lmax)
                nmin = min(lmin)
                w3.setRange(xRange=[idx - historyLength, idx], yRange=[nmin - 10, nmax + 10], padding=0)
    elif dispState == 1:
        if idx > 0:
            for n in range(0, 8):
                #labelCurve[n].setText(str(data[idx - 1, n]))
                if isDisplay[n] is True:
                    curve[n].setData(np.arange(0, idx), data[0:idx, n])
                    lmax[n] = max(data[0:idx, n])
                    lmin[n] = min(data[0:idx, n])
                else:
                    curve[n].setData([0])
            nmax = max(lmax)
            nmin = min(lmin)
            w3.setRange(xRange=[0, idx], yRange=[nmin, nmax + 10], padding=0)

    else:
        pass


def save():
    global state
    state = area.saveState()


def load():
    global state
    area.restoreState(state)


def serial_find():
    plist = list(serial.tools.list_ports.comports())
    if len(plist) <= 0:
        print("The Serial port can't find!")
    else:
        for plist_ch in plist:
            lss = list(plist_ch)
            cmbPortSelect.addItem(lss[0])
        # serialName = plist_0[0]
        # serialFd = serial.Serial(serialName, 115200, timeout=60)

        # print("check which port was really used >", serialFd.name)


def btnStartSampleClicked():
    global ser1
    global idx
    global dispState
    if ser1 is None:
        idx = 0
        dispState = 0
        Com = cmbPortSelect.currentText()
        BaudRate = 115200
        Parity = serial.PARITY_NONE
        DataBits = serial.EIGHTBITS
        StopBits = serial.STOPBITS_ONE

        ser1 = MySerialClass(Com=Com, BaudRate=BaudRate, Parity=Parity, DataBits=DataBits, StopBits=StopBits)

        btnSerOpen.setText('????????????')
        cmdStartSample.setText('  ????????????  ')
        global timer
        timer = pg.QtCore.QTimer()
        timer.timeout.connect(frameHandle)  # ????????????plotData??????
        timer.start(1)  # ??????ms????????????
    else:
        ser1.__del__()
        ser1 = None
        btnSerOpen.setText('????????????')
        cmdStartSample.setText('  ????????????  ')
        dispState = 1
        dataDisplay()


def cmdCurveClick(nn):
    global isDisplay
    if btnCurveDisp[nn].isChecked():
        isDisplay[nn] = True
    else:
        isDisplay[nn] = False
    dataDisplay()

def cmdSaveClick():
    savefiles()


def cmdLoadClick():
    openfiles()


def savefiles():
    global data
    global win
    global idx
    fname, type = QFileDialog.getSaveFileName(win, '????????????', "/", 'npy(*.npy)')
    # print(fname)
    # np.savetxt(fname, data)
    np.save(fname, data[0:idx])
    '''
    dlg.setFileMode(QFileDialog.AnyFile)
    dlg.setFilter(QDir.Files)
    if dlg.exec_():
        filenames = dlg.selectedFiles()
        np.savetxt(filenames, data)
    '''


def openfiles():
    global data
    global win
    global idx
    global dispState
    fname, type = QFileDialog.getOpenFileName(win, '????????????', "/", 'npy(*.npy)')
    data = np.load(fname)
    idx = len(data)
    dispState = 1
    dataDisplay()




if __name__ == '__main__':

    data = np.zeros((1000000, 8))
    historyLength = 1000  # ???????????????
    ser1 = None
    timer = None
    itemListNum = 8
    itemListNumY = 8
    isDisplay = [False, False, False, False, False, False, False, False]
    listSignalName = ['????????????', '????????????2', '????????????1', '????????????', '???????????????', '????????????', '????????????', '????????????']
    btnCurveColor = ['#ff0000', '#00ff00', '#0000ff', '#ffff00', '#ff00ff', '#00ffff', '#80ff80', '#808080']
    listCurveColor = [(255,0,0), (0,255,0), (0,0,255), (255,255,0), (255,0,255), (0,255,255), (80,255,80), (128,128,128)]
    btnCurveDisp = ['', '', '', '', '', '', '', '']
    listSignalPara = ['?????????', '?????????', '??????', '?????????', '?????????a', '?????????b', '?????????c']
    listTableItem = [['' for i in range(0,itemListNum)] for i in range(0, itemListNumY)]

    curve = ['', '', '', '', '', '', '', '']
    idx = 0
    dispState = 0
    state = None

    app = QtGui.QApplication([])
    win = QtGui.QMainWindow()
    area = DockArea()
    win.setCentralWidget(area)
    win.resize(800, 600)
    win.setWindowTitle('???????????????????????????V1.0')



    # Create docks, place them into the window one at a time.
    # Note that size arguments are only a suggestion; docks will still have to
    # fill the entire dock area and obey the limits of their internal widgets.
    d1 = Dock("Dock1", size=(1, 1))  # give this dock the minimum possible size
    d2 = Dock("Dock2 - Console", size=(500,50), closable=True)
    d3 = Dock("Dock3", size=(500, 300))
    d4 = Dock("Dock4 (tabbed) - Plot", size=(500, 200))
    d5 = Dock("Dock5 - Image", size=(500, 200))
    d6 = Dock("Dock6 (tabbed) - Plot", size=(500, 200))
    area.addDock(d3, 'top')
    area.addDock(d1, 'right', d3)
    area.addDock(d2, 'bottom')
    area.addDock(d4, 'bottom', d1)  ## place d4 at right edge of dock area
    area.addDock(d5, 'left')  ## place d5 at left edge of d1
    d5.hideTitleBar()
    # area.addDock(d6, 'top', d4)  ## place d5 at top edge of d4

    # area.moveDock(d4, 'top', d2)  ## move d4 to top edge of d2
    # area.moveDock(d6, 'above', d4)  ## move d6 to stack on top of d4
    # area.moveDock(d5, 'top', d2)  ## move d5 to top edge of d2

    # init dock1 layout
    d1.hideTitleBar()
    w1 = QtWidgets.QWidget()

    labelLogo = QtWidgets.QLabel()
    pm = QPixmap("./img/poc.jpg")
    scaredPixmap = pm.scaled(100, 80, aspectRatioMode=Qt.KeepAspectRatio)
    labelLogo.setPixmap(scaredPixmap)
    labelLogo.setFixedSize(100, 80)

    cmdStartSample = QtWidgets.QPushButton('   ????????????   ')
    cmdStartSample.clicked.connect(btnStartSampleClicked)
    cmdStartSample.setFixedSize(100, 30)
    cmdStartSample.setCheckable(True)
    cmdStartSample.setStyleSheet('''QPushButton{background:#00ff00;border-radius:5px;}
                                    QPushButton::checked{background:#ff0000;border-radius:5px;}''')

    cmdSaveData = QtWidgets.QPushButton('   ????????????   ')
    cmdSaveData.clicked.connect(cmdSaveClick)
    cmdSaveData.setFixedSize(100, 30)

    cmdLoadData = QtWidgets.QPushButton('   ????????????   ')
    cmdLoadData.clicked.connect(cmdLoadClick)
    cmdLoadData.setFixedSize(100, 30)

    cmbPortSelect = QtWidgets.QComboBox()
    cmbPortSelect.setFixedSize(100, 30)

    btnSerOpen = QtWidgets.QPushButton('????????????')
    btnSerOpen.clicked.connect(btnStartSampleClicked)

    groupCurveFitting = QtWidgets.QGroupBox('???????????????')
    layoutCurveFitting = QtWidgets.QHBoxLayout()
    labelFittingStyle = QtWidgets.QLabel("????????????")
    cmbFittingStyle = QtWidgets.QComboBox()
    txtFittingParaA = QtWidgets.QTextEdit('a=')
    txtFittingParaB = QtWidgets.QTextEdit('b=')
    txtFittingParaC = QtWidgets.QTextEdit('c=')
    layoutCurveFitting.addWidget(labelFittingStyle)
    groupCurveFitting.setLayout(layoutCurveFitting)

    layout1 = QtWidgets.QVBoxLayout()
    layout1.addWidget(labelLogo)
    layout1.addWidget(cmbPortSelect)
    layout1.addWidget(cmdStartSample)
    layout1.addWidget(cmdSaveData)

#    space = QSpacerItem(10, 100, QSizePolicy.Expanding, QSizePolicy.Minimum)
#    layout1.addItem(space)
    layout1.addWidget(cmdLoadData)
    w1.setLayout(layout1)
    d1.addWidget(w1)

    # init dock2 layout
    d2.hideTitleBar()
    tableWidget = QtWidgets.QTableWidget(8, 7)
    tableWidget.setVerticalHeaderLabels(listSignalName)
    tableWidget.setHorizontalHeaderLabels(listSignalPara)
    for m in range(8):
        for n in range(8):
            listTableItem[m][n] = QtWidgets.QLabel('')
            tableWidget.setCellWidget(m, n, listTableItem[m][n])
    btnDataJustPara = QtWidgets.QPushButton('??????????????????')
    btnDataJustPara.setFixedSize(100, 30)
    btnDataAnalysis = QtWidgets.QPushButton('????????????')
    btnDataAnalysis.setFixedSize(100, 30)
    btnReport = QtWidgets.QPushButton('??????????????????')
    btnReport.setFixedSize(100, 30)

    d2.addWidget(tableWidget, 0, 0, 8, 1)
    d2.addWidget(btnDataJustPara, 0, 1, 1, 1)
    d2.addWidget(btnDataAnalysis, 1, 1, 1, 1)
    d2.addWidget(btnReport, 2, 1, 1, 1)


    # init dock3 layout
    d3.hideTitleBar()

    w3 = pg.PlotWidget(title="??????")
    w3.resize(200, 100)
    # w3.plot(np.random.normal(size=100))
    # p = win.addPlot()#??????p??????????????????
    w3.showGrid(x=True, y=True)  # ???X???Y???????????????
    w3.setRange(xRange=[0, historyLength], yRange=[0, 10], padding=0)
    brush = QtGui.QBrush(QtGui.QColor(250, 250, 250))
    w3.setBackground(brush)
    w3.setLabel(axis='left', text='y / ?????????')  # ??????
    w3.setLabel(axis='bottom', text='x / ?????????')

    for n in range(0,itemListNum):
        curve[n] =w3.plot(pen = listCurveColor[n], name='curve')



    for i in range(len(btnCurveDisp)):
        btnCurveDisp[i] = QtWidgets.QPushButton(listSignalName[i])
        btnCurveDisp[i].clicked.connect(partial(cmdCurveClick, i))
        btnCurveDisp[i].setMinimumSize(100, 30)
        btnCurveDisp[i].setCheckable(True)
        btnCurveDisp[i].setStyleSheet('QPushButton::checked{{background:{0};border-radius:5px;}}'.format(btnCurveColor[i]))
    d3.addWidget(w3, 0, 0, 1, 8)

    d3.addWidget(btnCurveDisp[0], 9, 0, 1, 1)
    d3.addWidget(btnCurveDisp[1], 9, 1, 1, 1)
    d3.addWidget(btnCurveDisp[2], 9, 2, 1, 1)
    d3.addWidget(btnCurveDisp[3], 9, 3, 1, 1)
    d3.addWidget(btnCurveDisp[4], 9, 4, 1, 1)
    d3.addWidget(btnCurveDisp[5], 9, 5, 1, 1)
    d3.addWidget(btnCurveDisp[6], 9, 6, 1, 1)
    d3.addWidget(btnCurveDisp[7], 9, 7, 1, 1)

    '''
        cmdLoadData.setStyleSheet(''QPushButton{background:#000000;color:#ff0000;border-radius:5px;border-color:red
        ;border-width:20px;padding:5px;}QPushButton::pressed{background:#0000ff;color:#00ff00;}QToolButton::checked{
        background: #3C79F2;
        border-color: #11505C;
        font-weight: bold;
        font-family:"Microsoft YaHei";
        }'')
    '''
    '''
    labelCurve = ['','','','','','','','']
    for n in range(0,8):
        labelCurve[n] = QtWidgets.QLabel('')
        labelCurve[n].setFixedHeight(30)
        labelCurve[n].setAlignment(Qt.AlignCenter)

    d3.addWidget(labelCurve[0], 10, 0, 1, 1)
    d3.addWidget(labelCurve[1], 10, 1, 1, 1)
    d3.addWidget(labelCurve[2], 10, 2, 1, 1)
    d3.addWidget(labelCurve[3], 10, 3, 1, 1)
    d3.addWidget(labelCurve[4], 10, 4, 1, 1)
    d3.addWidget(labelCurve[5], 10, 5, 1, 1)
    d3.addWidget(labelCurve[6], 10, 6, 1, 1)
    d3.addWidget(labelCurve[7], 10, 7, 1, 1)
    '''
    # init dock4 layout
    #w4 = pg.PlotWidget(title="Dock 4 plot")
    #w4.plot(np.random.normal(size=100))
    flukeWin = MyFlukeWidget.MyFlukeWidget()
    d4.addWidget(flukeWin)
    d4.hideTitleBar()

    # init dock5 layout
    w5 = MyTestItemTree.MyTestItemTree()
    d5.addWidget(w5)

    # init dock6 layout
    w6 = pg.PlotWidget(title="Dock 6 plot")
    w6.plot(np.random.normal(size=100))
    d6.addWidget(w6)


    win.show()
    serial_find()



    #doc = docx.Document('.\test.docx')  # ????????????
    doc_new = docx.Document()  # ????????????



    doc_new.add_paragraph(u'?????????????????????????????????', style=None)
    doc_new.add_paragraph(u'?????????????????????????????????', style='Heading 2')

    # ???????????????????????????
    q = doc_new.paragraphs[1]
    q.add_run(u'???????????????????????????', style=None)

    #doc.save()  # ??????
    doc_new.save('./1.docx')  # ??????


    if (sys.flags.interactive != 1) or not hasattr(QtCore, 'PYQT_VERSION'):
        QtGui.QApplication.instance().exec_()
